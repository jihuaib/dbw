"""资料导入：Word / Markdown / 纯文本 → 命令清单。

刻意做得很轻：知识库只需要回答一件事 ——
**「这台设备上有哪些只读命令可用，各自看什么」**。
命令回显长什么样、怎么解析成结构化字段，全部交给 AI；
我们不再自建解析器体系。

两条提取路径：
  rule —— 正则扫 display/show 命令行，零依赖、可复现
  ai   —— 让模型通读资料后给出命令清单（结果按文档哈希缓存冻结）
"""
from __future__ import annotations

import html
import re
from typing import Any, Dict, List, Optional

from . import syntax as cli_syntax


# 已入库文档带这个版本。解析规则升级后，相同 SHA 的文档会自动重建命令，
# 不会继续被“内容重复”短路而永远保留旧的错误结果。
IMPORTER_VERSION = "KB-2.0.0"

# 只收只读命令。写命令永不进清单 —— 诊断不该改设备状态。
READ_VERBS = ("display", "show", "dis")
DANGEROUS = ("reset", "undo", "reboot", "delete", "erase", "save", "shutdown",
             "format", "install", "upgrade", "restore", "copy", "rename")

# 手册里命令行前面常见的「装饰」：标题井号、列表符、编号、标签（【命令】/命令格式：/语法：/Syntax:）
LINE_PREFIX_RE = re.compile(
    r"^\s*(?:#{1,6}\s*|[-*•]\s+|\d+(?:\.\d+)*\.?\s+)?"
    r"(?:【[^】]{1,12}】|(?:命令格式|命令|语法|用法|格式|Syntax|Command|Usage)\s*[:：])?\s*")
# 厂商语法行：display/show 开头，允许 [ ] { } | < > 与大小写参数名，整行取到末尾
SYNTAX_LINE_RE = re.compile(
    r"^(?:<[^>]{1,40}>|\[[^\]]{1,40}\])?\s*"
    r"((?:display|show|dis)\s+[A-Za-z0-9][A-Za-z0-9\-_./:{}\[\]|<>\s,*]{0,160}?)\s*[。；;]?\s*$")
# H3C/华为风格的必填参数不带 <>，靠命名习惯识别：vlan-id / interface-number / ip-address …
PARAM_WORD_RE = re.compile(
    r"^[a-z]+(?:-[a-z]+)*-(?:id|number|name|type|address|list|value|index|prefix|"
    r"mask|length|string|text|count|time|interval|ip|port)$")
# Markdown 表格行： | `show lldp neighbors` | global | 显示 LLDP 邻居概要 |
TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$", re.M)
BACKTICK_RE = re.compile(r"`([^`]+)`")
HEADING_RE = re.compile(r"^#{1,6}\s+(.+?)\s*$")
LABEL_RE = re.compile(r"^\s*【([^】]+)】\s*$")
READ_START_RE = re.compile(r"^(?:display|show|dis)\s+", re.I)
TABLE_HEADER_COMMANDS = {"命令", "命令格式", "command", "cli command", "syntax"}
TABLE_HEADER_VIEWS = {"视图", "view", "mode"}
TABLE_HEADER_PURPOSES = {"说明", "描述", "功能", "用途", "purpose", "description"}


def read_text(filename: str, raw: bytes) -> str:
    """Word / Markdown / 纯文本 → 统一纯文本。提取必须确定性。"""
    low = filename.lower()
    if low.endswith(".docx"):
        from io import BytesIO
        import docx
        doc = docx.Document(BytesIO(raw))
        parts: List[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    if low.endswith(".doc"):
        raise ValueError("旧版 .doc 不支持，请另存为 .docx")
    return raw.decode("utf-8", errors="replace")


def _clean(cmd: str) -> str:
    """只规整空白，**保留大小写** —— 不少设备的接口名是区分大小写的，
    统一转小写会让命令直接被拒。"""
    return re.sub(r"\s+", " ", cmd.strip())


def _key(cmd: str) -> str:
    """去重与匹配用的小写键。发送时仍用原始大小写。"""
    return _clean(cmd).lower()


def _excluded_commands() -> set:
    """观测者自照镜子的命令（CLI 会话 / 审计类）：输出必然带着采集自身的痕迹，
    进了证据快照就永不一致。清单是设置项（按厂商增删），不写死在代码里。"""
    from ..settings import service as settings
    raw = settings.get("kb_exclude_commands") or ""
    return {_key(x) for x in raw.split(",") if x.strip()}


def _is_read_only(cmd: str) -> bool:
    tokens = _key(cmd).split()
    if not tokens or tokens[0] not in READ_VERBS:
        return False
    if _key(cmd) in _excluded_commands():
        return False
    return not any(d in tokens for d in DANGEROUS)


def split_syntax(syntax: str) -> Dict[str, Any]:
    """把语法串拆成「可下发基串 + 必需参数」。

    `[]` 是可选组，`{a|b}` 是必选分支，`<param>` 是参数。嵌套结构和
    “参数后还有固定词”的语法统一由 :mod:`kb.syntax` 解析，导入与下发共用同一口径。
    """
    parsed = cli_syntax.analyze(syntax)
    return {"base": parsed["base"], "required": parsed["required"]}


def base_of_syntax(syntax: str) -> str:
    return split_syntax(syntax)["base"]


def _entry(syntax: str, purpose: str = "",
           keywords: Optional[List[str]] = None,
           sample: str = "") -> Optional[Dict[str, Any]]:
    """Build one command row from a complete syntax variant."""
    syntax = _clean(syntax)
    if not syntax or len(syntax) > 4096:
        return None
    try:
        parsed = cli_syntax.analyze(syntax)
    except ValueError:
        return None
    command = parsed["command"]
    literals = [str(token).lower() for token in parsed["literals"]]
    if any(token in DANGEROUS for token in literals) \
            or not cli_syntax.starts_with_one_of(syntax, READ_VERBS) \
            or not _is_read_only(command) or len(command.split()) < 2:
        return None
    words = {token for token in literals if token not in READ_VERBS}
    words.update(str(t).lower() for t in (keywords or []) if str(t).strip())
    return {
        "command": command,
        "syntax": syntax,
        "required": parsed["required"],
        "purpose": (purpose or "")[:200],
        "keywords": sorted(words),
        "params": sorted(set(parsed["params"])),
        "sample": sample,
        "read_only": True,
    }


def _plain_markdown(value: str, parameters: bool = False) -> str:
    """Unwrap Huawei-style Markdown while retaining CLI grammar punctuation."""
    # Remove real HTML tags before unescaping entities: ``&lt;1-32&gt;`` is a
    # CLI grammar annotation and must not be mistaken for an HTML tag.
    value = re.sub(
        r"</?(?:p|br|span|div|code|strong|em|table|tbody|thead|tr|th|td|ul|ol|li)"
        r"(?:\s+[^>]*)?>", " ", value or "", flags=re.I)
    value = html.unescape(value)
    value = re.sub(r"\s+", " ", value)
    value = re.sub(r"\\([\[\]{}|#])", r"\1", value)
    value = re.sub(r"\*\*(.+?)\*\*", r"\1", value)
    if parameters:
        def mark(match: re.Match) -> str:
            # Some manuals italicize two adjacent parameters as one span.
            names = match.group(1).split()
            return " ".join("<{0}>".format(name) for name in names)
        value = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", mark, value)
        # Canonicalize Huawei's repeat suffix as one grammar token.  The CLI
        # matcher interprets ``<community>&<1-32>`` as 1..32 values.
        value = re.sub(
            r"(<[^<>]+>)\s*&+\s*<(\d+(?:-\d+)?)>", r"\1&<\2>", value)
    else:
        value = value.replace("*", "")
    value = re.sub(r"[\u00ad\u200b\u200c\u200d\ufeff]", "", value)
    return _clean(value)


def extract_by_command_blocks(text: str) -> List[Dict[str, Any]]:
    """Extract formal variants from independent ``【命令】`` sections.

    Large Huawei manuals wrap fixed words in ``**bold**``, parameters in
    ``*italic*``, escape square brackets, and fold one syntax across many
    physical lines.  Paragraph boundaries, not individual lines, delimit the
    variants.  Once this high-confidence structure is present it is safer than
    scanning examples and explanatory tables elsewhere in the document.
    """
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    found: Dict[str, Dict[str, Any]] = {}
    heading = ""
    intro: List[str] = []
    i = 0
    while i < len(lines):
        title = HEADING_RE.match(lines[i])
        if title:
            heading = _plain_markdown(title.group(1))
            intro = []
            i += 1
            continue
        label = LABEL_RE.match(lines[i])
        if label and label.group(1).strip() == "命令":
            i += 1
            body: List[str] = []
            while i < len(lines) and not LABEL_RE.match(lines[i]) \
                    and not HEADING_RE.match(lines[i]):
                body.append(lines[i])
                i += 1
            purpose = _plain_markdown("\n".join(intro))
            paragraphs = re.split(
                r"\n[\t \u00a0\u200b\u200c\ufeff]*\n+",
                "\n".join(body).strip())
            heading_words = [t for t in heading.split()
                             if t.lower() not in READ_VERBS]
            for paragraph in paragraphs:
                syntax = _plain_markdown(paragraph, parameters=True)
                if not READ_START_RE.match(syntax):
                    continue
                entry = _entry(syntax, purpose, heading_words)
                if entry:
                    found[_key(syntax)] = entry
            continue
        if heading and not label:
            intro.append(lines[i])
        i += 1
    return [found[k] for k in sorted(found)]


def _split_markdown_cells(body: str) -> List[str]:
    """Split a Markdown row without cutting escaped/code-span ``|`` tokens."""
    cells: List[str] = []
    buf: List[str] = []
    in_code = False
    i = 0
    while i < len(body):
        ch = body[i]
        if ch == '\\' and i + 1 < len(body) and body[i + 1] == '|':
            buf.append("|")
            i += 2
            continue
        if ch == "`":
            in_code = not in_code
            buf.append(ch)
        elif ch == "|" and not in_code:
            cells.append("".join(buf).strip())
            buf = []
        else:
            buf.append(ch)
        i += 1
    cells.append("".join(buf).strip())
    return cells


def _table_header(value: str) -> str:
    value = BACKTICK_RE.sub(lambda match: match.group(1), value)
    return re.sub(r"[*_\s]+", " ", value).strip().lower()


def _separator_cell(value: str) -> bool:
    return bool(value) and set(value) <= set("-: ")


def extract_by_markdown_table(text: str) -> List[Dict[str, Any]]:
    """解析 Markdown 表格式 CLI 文档。

    形如「| 命令 | 视图 | 说明 |」的 Markdown 表格手册：
        | 命令 | 视图 | 说明 |
        | `show lldp neighbors` | global | 显示 LLDP 邻居概要 |

    这类文档只给命令与说明、不给样例回显 —— 正好，回显怎么读交给 AI。
    """
    found: Dict[str, Dict[str, Any]] = {}
    section = ""
    command_col, view_col, purpose_col = 0, 1, 2
    for line in text.replace("\r\n", "\n").split("\n"):
        stripped = line.strip()
        if stripped.startswith("#"):
            section = stripped.lstrip("#").strip()
            continue
        m = TABLE_ROW_RE.match(line)
        if not m:
            continue
        cells = _split_markdown_cells(m.group(1))
        if len(cells) < 2 or all(_separator_cell(c) for c in cells):
            continue
        headers = [_table_header(c) for c in cells]
        command_headers = [i for i, value in enumerate(headers)
                           if value in TABLE_HEADER_COMMANDS]
        if command_headers:
            command_col = command_headers[0]
            view_col = next((i for i, value in enumerate(headers)
                             if value in TABLE_HEADER_VIEWS), -1)
            purpose_col = next((i for i, value in enumerate(headers)
                                if value in TABLE_HEADER_PURPOSES), -1)
            continue
        if command_col >= len(cells):
            continue
        marks = BACKTICK_RE.findall(cells[command_col])
        syntax = (marks[0] if marks else cells[command_col]).strip()
        syntax = _plain_markdown(syntax, parameters=True)
        if not READ_START_RE.match(syntax):
            continue
        view = cells[view_col] if 0 <= view_col < len(cells) else ""
        purpose = cells[purpose_col] if 0 <= purpose_col < len(cells) else ""
        entry = _entry(syntax, purpose or section, [view] if view else [])
        if entry:
            found[_key(syntax)] = entry
    return sorted(found.values(), key=lambda item: (
        _key(item["command"]), _key(item["syntax"])))


def looks_like_table_doc(text: str) -> bool:
    # 独立【命令】块比说明表格更权威；厂商手册里的大量字段表不能把整份
    # 文档误分流到 table 引擎。
    for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        label = LABEL_RE.match(line)
        if label and label.group(1).strip() == "命令":
            return False
    rows = TABLE_ROW_RE.findall(text)
    if len(rows) < 2:
        return False
    command_col = 0
    saw_header = False
    for raw in rows:
        cells = _split_markdown_cells(raw)
        headers = [_table_header(c) for c in cells]
        indexes = [i for i, value in enumerate(headers)
                   if value in TABLE_HEADER_COMMANDS]
        if indexes:
            command_col, saw_header = indexes[0], True
            continue
        if all(_separator_cell(c) for c in cells) or command_col >= len(cells):
            continue
        cell = cells[command_col]
        marks = BACKTICK_RE.findall(cell)
        syntax = _plain_markdown(marks[0] if marks else cell, parameters=True)
        if READ_START_RE.match(syntax):
            return saw_header or len(rows) >= 3
    return False


def extract_by_inline(text: str) -> List[Dict[str, Any]]:
    """提取行内反引号里的命令 —— 覆盖「标题/列表」式文档。

    形如「### `命令`」标题 + 「- **用法**：`命令`」列表的手册：
        ### 2.3 `show current-configuration`
        - **用法**：`show current-configuration`

    这类文档不给表格也不给裸命令行，命令只出现在 `反引号` 里。
    漏掉它的代价在实测中出现过：`show current-configuration` 没进清单，
    闭集校验让 Agent 无权采配置，shutdown 的接口只能诊断成「物理链路未建立」。
    """
    found: Dict[str, Dict[str, Any]] = {}
    lines = text.replace("\r\n", "\n").split("\n")
    section = ""
    desc = ""
    await_desc: List[str] = []   # 等着用标题下第一行正文回填 purpose 的命令
    for line in lines:
        stripped = line.strip()
        is_heading = stripped.startswith("#")
        if is_heading:
            section = re.sub(r"^[#\d\.\s]+", "", stripped).strip()
            desc = ""
            await_desc = []
        elif stripped and not stripped.startswith(("-", "|", "`", ">")):
            # 标题后的第一行正文当作该节命令的用途说明
            if not desc:
                desc = stripped[:120]
                for k in await_desc:
                    if k in found and not found[k]["purpose"]:
                        found[k]["purpose"] = desc
                await_desc = []
        for mark in BACKTICK_RE.findall(line):
            syntax = _plain_markdown(mark, parameters=True)
            entry = _entry(syntax)
            if not entry:
                continue
            base = entry["command"]
            identity = _key(syntax)
            if identity in found:
                continue
            purpose = desc or BACKTICK_RE.sub(
                lambda m: m.group(1), section).strip()[:120]
            if _key(purpose) == _key(base):
                purpose = ""
            entry["purpose"] = purpose
            found[identity] = entry
            if is_heading and not purpose:
                await_desc.append(identity)
    return sorted(found.values(), key=lambda item: (
        _key(item["command"]), _key(item["syntax"])))


def _normalize_syntax_line(line: str) -> str:
    """去掉标题/列表/编号/标签装饰，留下可能是命令的正文。"""
    return LINE_PREFIX_RE.sub("", line, count=1).strip()


def _mark_bare_params(syntax: str) -> str:
    """把不带 <> 的必填参数（按命名习惯识别）包成 <param>，交给 split_syntax。"""
    out = []
    for i, tok in enumerate(syntax.split()):
        bare = tok.strip("[]{}|,")
        if i >= 1 and PARAM_WORD_RE.match(bare) and "<" not in tok:
            tok = tok.replace(bare, "<" + bare + ">")
        out.append(tok)
    return " ".join(out)


def extract_by_rule(text: str) -> List[Dict[str, Any]]:
    """正则提取：示例行 / 标题行 / 语法行 / 代码块 / 行内反引号，多路合并。

    覆盖的写法（各厂商手册常见）：
      <Sysname> display ospf peer                      示例行（带提示符）
      # display ospf peer  /  ## 1.2 display interface  标题就是命令
      【命令】 display vlan vlan-id                     标签 + 命令
      display ospf [ process-id ] peer [ verbose ]      「命令格式」段的语法行
      ```display ip routing-table [ verbose ]```        代码块
    同一完整语法多次出现只留一条（保留样例回显最长的），按命令串排序。
    """
    structured = extract_by_command_blocks(text)
    if structured:
        # 命令块是厂商手册的正式定义。继续扫全文会把示例命令、字段说明和
        # “1. display ... statistics 命令显示信息描述表”一类标题误收入库。
        return structured

    found: Dict[str, Dict[str, Any]] = {}
    lines = text.replace("\r\n", "\n").split("\n")
    for idx, line in enumerate(lines):
        body = _normalize_syntax_line(line)
        if not body or body.startswith("|"):
            continue
        m = SYNTAX_LINE_RE.match(body)
        if not m:
            continue
        syntax = _mark_bare_params(_clean(m.group(1)))
        entry = _entry(syntax)
        if not entry:
            continue
        cmd = entry["command"]
        # 命令行之后连续的非空行当作样例回显（示例行才有；语法行后面通常是说明）
        sample: List[str] = []
        if line.strip().startswith("<") or line.strip().startswith("["):
            j = idx + 1
            while j < len(lines) and lines[j].strip() and not SYNTAX_LINE_RE.match(
                    _normalize_syntax_line(lines[j])):
                sample.append(lines[j])
                if len(sample) >= 24:
                    break
                j += 1
        # 命令行之前最近的非空行当作用途说明
        purpose = ""
        k = idx - 1
        while k >= 0 and k > idx - 6:
            prev = lines[k].strip()
            if prev and not SYNTAX_LINE_RE.match(_normalize_syntax_line(prev)) \
                    and not prev.startswith("```"):
                purpose = LINE_PREFIX_RE.sub("", prev, count=1).strip()[:120]
                break
            k -= 1
        identity = _key(syntax)
        prev_entry = found.get(identity)
        if prev_entry:
            # 同一语法重复出现时保留更完整的样例。
            if len("\n".join(sample)) > len(prev_entry["sample"]):
                prev_entry["sample"] = "\n".join(sample)
            continue
        entry["purpose"] = purpose
        entry["sample"] = "\n".join(sample)
        found[identity] = entry
    for c in extract_by_inline(text):
        identity = _key(c["syntax"])
        if identity not in found:   # 裸命令行带样例回显，优先
            found[identity] = c
    return sorted(found.values(), key=lambda item: (
        _key(item["command"]), _key(item["syntax"])))


AI_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "properties": {
        "commands": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "command": {"type": "string",
                                "description": "完整只读命令，小写，不含设备提示符"},
                    "purpose": {"type": "string", "description": "一句话说明它看什么"},
                    "keywords": {"type": "array", "items": {"type": "string"},
                                 "description": "该命令相关的中英文检索词"},
                    "params": {"type": "array", "items": {"type": "string"},
                               "description": "可变参数名，如 ip-address / interface"},
                },
                "required": ["command", "purpose", "keywords", "params"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["commands"],
    "additionalProperties": False,
}

AI_SYSTEM = (
    "你是网络设备 CLI 资料的整理助手。从给定资料中列出所有**只读诊断命令**"
    "（display / show 开头），忽略任何会改变设备状态的命令。"
    "只依据原文，不得臆造命令。keywords 要包含中文口语说法（如「路由表」「邻居」「光模块」），"
    "方便后续按用户提问检索。"
)


def extract_by_ai(text: str) -> Dict[str, Any]:
    from ...core import llm
    body = text if len(text) <= 60000 else text[:60000]
    res = llm.call_json("kb.extract", AI_SYSTEM, "资料原文：\n\n" + body, AI_SCHEMA)
    if not res["ok"]:
        return {"ok": False, "commands": [], "error": res["error"]}
    out: List[Dict[str, Any]] = []
    for c in res["data"].get("commands", []):
        entry = _entry(str(c.get("command", "")))
        if not entry:
            continue
        entry["purpose"] = str(c.get("purpose", ""))[:200]
        entry["keywords"] = sorted(
            set(entry["keywords"]) |
            {str(k).lower() for k in c.get("keywords", []) if k})
        entry["params"] = sorted(
            set(entry["params"]) |
            {str(p) for p in c.get("params", []) if p})
        out.append(entry)
    uniq = {_key(c["syntax"]): c for c in out}
    return {"ok": True, "commands": sorted(
                uniq.values(), key=lambda item: (
                    _key(item["command"]), _key(item["syntax"]))),
            "cached": res.get("cached", False), "error": ""}
